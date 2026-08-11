"""Tests for document extraction, chunking, embeddings/FAISS retrieval, and citations."""
import io

from src import rag


def test_chunk_text_short_text_single_chunk():
    text = "This is a short procurement policy statement."
    assert rag.chunk_text(text) == [text]


def test_chunk_text_overlapping_long_text():
    text = " ".join(f"word{i}" for i in range(500))
    chunks = rag.chunk_text(text, chunk_size=200, overlap=50)
    assert len(chunks) > 1
    for c in chunks:
        assert len(c) <= 220  # small slack for word-boundary snapping
    tail_words = chunks[0].split()[-3:]
    assert any(w in chunks[1] for w in tail_words), "consecutive chunks should overlap"


def test_extract_txt_pages():
    pages = rag.extract_txt_pages("Minimum delivery experience: 5 years.")
    assert pages == [(None, "Minimum delivery experience: 5 years.")]


def test_extract_txt_pages_empty_returns_nothing():
    assert rag.extract_txt_pages("   ") == []


def test_extract_docx_pages_roundtrip():
    import docx

    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Tender Requirement: vendors must have at least 5 years delivery experience.")
    document.add_paragraph("Compliance certifications are mandatory.")
    document.save(buf)
    buf.seek(0)

    pages = rag.extract_docx_pages(buf)
    assert len(pages) == 1
    page_num, text = pages[0]
    assert page_num is None
    assert "5 years delivery experience" in text
    assert "Compliance certifications" in text


def test_extract_pdf_pages_roundtrip():
    import fitz

    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page one: minimum delivery experience is 5 years.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page two: compliance certification required.")
    data = doc.tobytes()
    doc.close()

    pages = rag.extract_pdf_pages(data)
    assert len(pages) == 2
    assert pages[0][0] == 1
    assert "minimum delivery experience" in pages[0][1].lower()
    assert pages[1][0] == 2
    assert "compliance certification" in pages[1][1].lower()


def test_extract_document_pages_rejects_unsupported_type():
    import pytest

    with pytest.raises(ValueError):
        rag.extract_document_pages("spreadsheet.xlsx", io.BytesIO(b"whatever"))


def test_chunk_document_preserves_citation_metadata():
    pages = [(1, "Vendors must maintain at least five years of delivery experience in similar contracts.")]
    chunks = rag.chunk_document("doc123", "tender.pdf", pages)
    assert len(chunks) >= 1
    c = chunks[0]
    assert c.doc_id == "doc123"
    assert c.filename == "tender.pdf"
    assert c.page == 1
    assert c.chunk_id.startswith("doc123-")


def test_format_citation_includes_page_when_available():
    chunk = rag.DocumentChunk(chunk_id="d-1", doc_id="d", filename="tender.pdf", page=4, text="...")
    citation = rag.format_citation(chunk)
    assert "tender.pdf" in citation
    assert "p.4" in citation
    assert "d-1" in citation


def test_format_citation_omits_page_when_unavailable():
    chunk = rag.DocumentChunk(chunk_id="d-2", doc_id="d", filename="policy.docx", page=None, text="...")
    citation = rag.format_citation(chunk)
    assert "p." not in citation


def test_knowledge_base_retrieve_before_build_returns_empty():
    kb = rag.KnowledgeBase()
    assert kb.is_ready is False
    assert kb.retrieve("anything") == []


def test_knowledge_base_build_with_no_extractable_text():
    kb = rag.KnowledgeBase()
    report = kb.build([("empty.txt", [])])
    assert report["ok"] is False
    assert kb.is_ready is False


def test_knowledge_base_build_and_retrieve_relevance():
    documents = [
        (
            "delivery_policy.txt",
            [(None, "All vendors must demonstrate a minimum of five years of prior delivery experience on comparable contracts.")],
        ),
        ("kitchen_menu.txt", [(None, "Today's cafeteria menu includes grilled chicken, rice, and steamed vegetables.")]),
    ]
    kb = rag.KnowledgeBase()
    report = kb.build(documents)
    assert report["ok"] is True
    assert report["chunk_count"] >= 2
    assert kb.is_ready is True

    results = kb.retrieve("What is the minimum delivery experience required?", k=2)
    assert results
    assert results[0].chunk.filename == "delivery_policy.txt"
    scores_by_file = {r.chunk.filename: r.score for r in results}
    if "kitchen_menu.txt" in scores_by_file:
        assert scores_by_file["delivery_policy.txt"] > scores_by_file["kitchen_menu.txt"]


def test_knowledge_base_no_relevant_document_scenario():
    """An empty/unbuilt knowledge base must report no results, not fabricate a match."""
    kb = rag.KnowledgeBase()
    assert kb.retrieve("what is the payment schedule?") == []
